"""One-off Markdown -> .docx converter for the FYP report draft.

Not a general-purpose converter — handles exactly the markdown subset
used in FYP_Report.md: #/##/### headings, pipe tables, bold/italic/code
inline spans, bullet and numbered lists, blockquote-style italic notes,
and horizontal rules (rendered as a page break between major chapters).
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(__file__).parent / "FYP_Report.md"
OUT = Path(sys.argv[2]) if len(sys.argv) > 2 else SRC.with_suffix(".docx")

INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")


def add_inline_runs(paragraph, text):
    for chunk in INLINE_RE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("*") and chunk.endswith("*") and not chunk.startswith("**"):
            run = paragraph.add_run(chunk[1:-1])
            run.italic = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(chunk)


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def parse_table(lines, i):
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(row)
        i += 1
    # drop the "---|---" separator row (row index 1)
    if len(rows) > 1 and all(set(c) <= set("-: ") for c in rows[1]):
        rows.pop(1)
    return rows, i


def build(doc, lines):
    i = 0
    n = len(lines)
    first_h1 = True
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if stripped == "":
            i += 1
            continue

        if stripped == "---":
            # page break between major sections, skip on very first lines
            if doc.paragraphs:
                doc.add_page_break()
            i += 1
            continue

        if stripped.startswith("| "):
            rows, i = parse_table(lines, i)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r, row in enumerate(rows):
                    for c, cell_text in enumerate(row):
                        if c >= len(table.columns):
                            continue
                        cell = table.cell(r, c)
                        cell.text = ""
                        p = cell.paragraphs[0]
                        add_inline_runs(p, cell_text)
                        if r == 0:
                            for run in p.runs:
                                run.bold = True
                            set_cell_shading(cell, "D9E2F3")
                doc.add_paragraph("")
            continue

        m = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                if first_h1:
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.bold = True
                    run.font.size = Pt(20)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    first_h1 = False
                else:
                    doc.add_heading(text, level=1)
            elif level == 2:
                doc.add_heading(text, level=2)
            else:
                doc.add_heading(text, level=3)
            i += 1
            continue

        # bullet list
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:])
            i += 1
            continue

        # numbered list
        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, m.group(2))
            i += 1
            continue

        # bold-only short "field: value" cover-page lines
        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(p, stripped)
            i += 1
            continue

        # italic-only note/placeholder paragraph
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped[1:-1])
            run.italic = True
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            i += 1
            continue

        # normal paragraph
        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        p.paragraph_format.space_after = Pt(8)
        i += 1


def main():
    text = SRC.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for lvl, size in ((1, 16), (2, 13), (3, 12)):
        hstyle = doc.styles[f"Heading {lvl}"]
        hstyle.font.size = Pt(size)
        hstyle.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        hstyle.font.bold = True

    build(doc, lines)
    doc.save(str(OUT))
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
